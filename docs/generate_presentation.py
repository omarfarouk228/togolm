from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TOGO_GREEN = RGBColor(0x00, 0x6a, 0x4e)
TOGO_YELLOW = RGBColor(0xFF, 0xCE, 0x00)
TOGO_RED = RGBColor(0xD2, 0x10, 0x34)
DARK = RGBColor(0x1e, 0x29, 0x3b)
GRAY = RGBColor(0x64, 0x74, 0x8b)
LIGHT_GRAY = RGBColor(0xf1, 0xf5, 0xf9)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK


def set_font(run, bold=False, size=11, color=DARK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'


def heading(text, level=1, color=TOGO_GREEN, size=18, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, bold=True, size=size, color=color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '006a4e')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def stats_table(rows):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(['Métrique', 'Valeur']):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        set_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '006a4e')
        tcPr.append(shd)
    for idx, (metric, value) in enumerate(rows):
        row = table.rows[idx + 1].cells
        row[0].text = metric
        row[1].text = value
        for cell in row:
            set_font(cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(''))
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TogoLM')
set_font(r, bold=True, size=32, color=TOGO_GREEN)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Première infrastructure IA open-source pour le Togo')
set_font(r2, size=14, color=GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Omar Farouk KOUGBADA  ·  KOF CORPORATION  ·  Lomé, Togo')
set_font(r3, size=11, color=GRAY, italic=True)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. LE PROBLÈME
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Le problème', size=14)
body(
    'Les données publiques togolaises — lois, institutions, économie, presse — '
    'sont éparpillées et absentes des grands modèles d\'IA comme ChatGPT ou Gemini. '
    'Interroger un LLM sur le Togo donne des réponses génériques, imprécises ou vides.'
)
body(
    'Ce vide crée un désavantage pour les développeurs, entreprises et institutions '
    'togolaises qui veulent intégrer l\'IA dans leurs produits.'
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. LA SOLUTION
# ══════════════════════════════════════════════════════════════════════════════
heading('2. La solution — TogoLM', size=14)
body(
    'TogoLM est une infrastructure IA complète, open-source, construite autour '
    'de la connaissance togolaise. Elle repose sur quatre briques :'
)
bullet('Corpus structuré : 62 168 documents togolais collectés, nettoyés et vectorisés', bold_prefix='')
bullet('Moteur RAG : recherche vectorielle + Gemini 2.5 Flash avec tokens de réflexion', bold_prefix='')
bullet('Modèle fine-tuné : QLoRA sur Mistral 7B Instruct v0.3 — publié sur HuggingFace', bold_prefix='')
bullet('API publique REST : endpoints documentés, clé gratuite, déployée sur Railway', bold_prefix='')

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. CHIFFRES CLÉS
# ══════════════════════════════════════════════════════════════════════════════
heading('3. Chiffres clés (v1.0.1 — juin 2026)', size=14)
stats_table([
    ('Documents indexés',    '62 168'),
    ('Chunks vectorisés',    '310 840'),
    ('Dimensions embedding', '384 (gemini-embedding-001)'),
    ('Modèle RAG',           'Gemini 2.5 Flash (thinking_budget=2048)'),
    ('Modèle fine-tuné',     'Mistral 7B Instruct v0.3 — QLoRA (HuggingFace)'),
    ('API déployée',         'https://api.togolm.kofcorporation.com'),
    ('Statut',               'Production — live'),
])

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. STACK TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Stack technique', size=14)
bullet('Pipeline de collecte : Python · Scrapy · uv')
bullet('Stockage : PostgreSQL + pgvector (Supabase) · Alembic migrations')
bullet('Embeddings : gemini-embedding-001 (384-dim) · fallback MiniLM-L12-v2')
bullet('RAG : Gemini 2.5 Flash · SSE streaming · ThinkingBlock')
bullet('Fine-tuning : QLoRA · HuggingFace PEFT · SFTTrainer · Colab')
bullet('API : FastAPI · Pydantic · Railway (Nixpacks)')
bullet('Showcase : Next.js 16 · Tailwind v4')

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ENDPOINTS API
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Endpoints API disponibles', size=14)
bullet('GET  /v1/stats — statistiques live du corpus (public)')
bullet('GET  /v1/documents — liste paginée avec filtres source/catégorie/langue')
bullet('GET  /v1/search?q= — recherche plein-texte en français')
bullet('POST /v1/query — RAG : recherche vectorielle → Gemini 2.5 Flash → réponse complète')
bullet('POST /v1/query/stream — RAG streaming SSE (thinking · chunk · sources)')
bullet('POST /v1/embed — génération d\'embedding 384-dim')
bullet('POST /v1/auth/register — création de clé API gratuite (200 req/jour)')

p_note = doc.add_paragraph()
r_note = p_note.add_run('Documentation complète : https://api.togolm.kofcorporation.com/docs')
set_font(r_note, italic=True, color=GRAY, size=10)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PLANS API
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Plans API', size=14)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.LEFT
headers2 = ['Plan', 'Quota / jour', 'Accès']
hdr2 = table2.rows[0].cells
for i, h in enumerate(headers2):
    hdr2[i].text = h
    run = hdr2[i].paragraphs[0].runs[0]
    set_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    tc = hdr2[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '006a4e')
    tcPr.append(shd)
data2 = [
    ('Anonyme',      '20',       'Par IP'),
    ('Gratuit',      '200',      'Via /v1/auth/register'),
    ('Dev',          '1 000',    'Sur demande'),
    ('Institution',  '100 000',  'Partenaires'),
]
for idx, (plan, quota, acces) in enumerate(data2):
    row = table2.rows[idx + 1].cells
    row[0].text = plan
    row[1].text = quota
    row[2].text = acces
doc.add_paragraph()

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. BESOINS ACTUELS
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Besoins actuels', size=14)

heading('7.1  Visibilité et adoption', size=12, color=DARK)
body(
    'Le projet est techniquement complet et déployé en production. '
    'Le besoin prioritaire est d\'atteindre les développeurs et entreprises '
    'qui opèrent au Togo et en zone UEMOA, via des relais dans l\'écosystème tech africain.'
)

heading('7.2  Premier partenaire d\'intégration', size=12, color=DARK)
body(
    'Nous cherchons un partenaire business dont le produit est complémentaire à TogoLM '
    '— notamment des startups fintech, des plateformes de paiement ou des services '
    'opérant localement qui peuvent bénéficier du contexte réglementaire et institutionnel '
    'fourni par notre API.'
)

heading('7.3  Compute pour les prochaines itérations', size=12, color=DARK)
body(
    'Le fine-tuning actuel est publié, mais les itérations suivantes '
    '(plus de données, meilleur modèle de base) nécessitent un accès GPU fiable. '
    'Des crédits cloud (GCP, AWS, RunPod) ou un accès infrastructure seraient '
    'un accélérateur majeur.'
)

heading('7.4  Données supplémentaires', size=12, color=DARK)
body(
    'Des connexions avec des institutions togolaises ou ouest-africaines pouvant '
    'ouvrir l\'accès à des données structurées (MEF, togopress, republicoftogo...) '
    'permettraient d\'enrichir significativement le corpus.'
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
heading('8. Roadmap', size=14)
bullet('Crawl des sources restantes : togofirst, togoinfos, republicoftogo, togopress, mef')
bullet('Cross-encoder reranker pour améliorer la qualité de récupération RAG')
bullet('Extension du corpus à l\'ensemble de l\'Afrique de l\'Ouest francophone')
bullet('MCP Server TogoLM — intégration dans des écosystèmes d\'agents IA')
bullet('Soumission Show HN')

divider()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(20)
r_f1 = p_footer.add_run('API : https://api.togolm.kofcorporation.com  ·  ')
set_font(r_f1, color=GRAY, size=10)
r_f2 = p_footer.add_run('GitHub : github.com/omarfarouk228/togolm')
set_font(r_f2, color=GRAY, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), 'TogoLM_Presentation.docx')
doc.save(out)
print(f'Saved → {out}')
